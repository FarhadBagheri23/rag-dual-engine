"""Raw file -> text -> overlapping chunks.

Parsing is per-format; everything after it is format-agnostic, which is why
adding the URL scraper in phase 6 only means adding one `_parse_*` function.
"""

import re
import uuid
from pathlib import Path

import pymupdf
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.exceptions import EmptyDocument, UnsupportedFileType

SUPPORTED = {".pdf", ".docx"}


def _parse_pdf(path: Path) -> tuple[str, str]:
    """Returns (text, embedded title). The title is "" when the file carries no
    metadata title — naming the document is the caller's job, not this one's."""
    with pymupdf.open(path) as doc:
        text = "\n\n".join(page.get_text() for page in doc)
        title = (doc.metadata or {}).get("title") or ""
    return text, title.strip()


def _parse_docx(path: Path) -> tuple[str, str]:
    doc = DocxDocument(path)
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text, (doc.core_properties.title or "").strip()


PARSERS = {".pdf": _parse_pdf, ".docx": _parse_docx}

# Word-counted splitter: the spec specifies chunk sizes in words, not characters.
_splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.chunk_size,
    chunk_overlap=settings.chunk_overlap,
    length_function=lambda t: len(t.split()),
)


def parse(path: Path, filename: str) -> tuple[str, str]:
    """Extract text and a display title.

    `path` is where the bytes currently sit — a temp file during upload — and
    `filename` is what the user called the document. Keeping them separate is
    the whole point: falling back to `path.stem` is how a document ends up
    titled `tmpoqta_c1i`. Both the title fallback and the error message below
    use `filename`, the only one of the two a person ever chose.
    """
    suffix = path.suffix.lower()
    if suffix not in PARSERS:
        raise UnsupportedFileType(suffix, SUPPORTED)
    text, title = PARSERS[suffix](path)
    text = re.sub(r"[ \t]+", " ", text).strip()  # collapse the whitespace PDFs emit
    if not text:
        raise EmptyDocument(filename)
    return text, title or Path(filename).stem


def chunk(text: str, doc_id: str) -> list[dict]:
    return [
        {"id": f"{doc_id}:{i}", "ordinal": i, "text": piece}
        for i, piece in enumerate(_splitter.split_text(text))
    ]


def new_doc_id() -> str:
    return uuid.uuid4().hex[:12]
